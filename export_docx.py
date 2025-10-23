import sqlite3
from docx import Document

DB_FILE = "transcripts.db"

def export_all_transcripts(output_file="All_Transcripts.docx"):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()

    c.execute("SELECT audio_file, transcription, translation, language FROM transcripts")
    rows = c.fetchall()

    doc = Document()
    doc.add_heading("Call Transcriptions & Translations", level=0)

    for row in rows:
        audio_file, transcription, translation, language = row
        
        doc.add_heading(f"Audio File: {audio_file}", level=1)
        doc.add_paragraph(f"Language: {language if language else 'N/A'}")

        doc.add_heading("Transcription", level=2)
        doc.add_paragraph(transcription if transcription else "N/A")

        doc.add_heading("Translation", level=2)
        doc.add_paragraph(translation if translation else "N/A")

        doc.add_paragraph("\n")

    doc.save(output_file)
    print(f"✅ Export completed: {output_file}")
